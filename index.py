import pdfplumber
import json
import os
from openai import OpenAI
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- CONFIGURATION ---
OPENAI_API_KEY = "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"
# OPENROUTER_API_KEY = "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"

TEMPLATE_PATH = "chakli.docx"
OUTPUT_PATH = os.path.join("outputs", "Final_Generated_Resume.docx")
PDF_PATH = "test.pdf"

client = OpenAI(api_key=OPENAI_API_KEY)
# client = OpenAI(
#   base_url="https://openrouter.ai/api/v1",
#   api_key=OPENROUTER_API_KEY,
# )


# ==============================================================================
# TEMPLATE VARIABLE MAP  (verified from ravan.docx XML — do not change)
#
#  personal_info.name / .email / .phone / .location
#  summary
#  education[]          → e.degree, e.year
#  expertise_areas[]    → item  (plain string list)
#  certifications[]     → c     (plain string list)
#  awards[]             → a     (plain string list)
#  skills[]             → populated programmatically via populate_skills_table()
#
#  experience_summary[] → used by BOTH template sections with different aliases:
#
#    Section 1 — Experience Summary table  (alias: s)
#      s.role            → job title
#      s.years           → date range string
#      s.skills          → comma-separated skills string
#      s.achievements[]  → short bullet strings  (inner loop: bullet)
#
#    Section 2 — Professional Experience  (alias: exp)
#      exp.project_name         → named project/product/system
#      exp.organisation         → employer or client company
#      exp.duration             → date range string (same as s.years)
#      exp.role                 → job title (same as s.role)
#      exp.project_description  → 2-3 line project description
#      exp.key_activities[]     → responsibility bullets  (inner loop: activity)
#      exp.tools_used[]         → tool name strings       (joined with ', ')
#      exp.technologies_used[]  → tech/language strings   (joined with ', ')
#      exp.skills_used[]        → skill name strings      (joined with ', ')
#      exp.major_achievements[] → quantified achievement bullets (inner loop: achievement)
#      exp.client_appreciation[]→ client quote strings    (inner loop: appreciation)
#
#  Every experience_summary object must carry ALL of the above keys because
#  the template loops over the same array twice with different aliases.
# ==============================================================================


# ==============================================================================
# STEP 1: PDF EXTRACTION
# ==============================================================================

def extract_text_from_pdf(pdf_path, status_callback=None):
    """Reads every page of the PDF and returns the complete raw text."""
    msg = f"Reading PDF: {pdf_path}..."
    print(msg)
    if status_callback: status_callback(msg)
    
    text = "" 
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text


# ==============================================================================
# STEP 2: AI EXTRACTION — TWO-PASS APPROACH
#
# Pass 1 — Extract EVERYTHING verbatim. No filtering, no loss.
# Pass 2 — Reformat into the exact JSON shape the template needs.
#
# Two passes prevent the AI from silently dropping CSS, Git, soft skills,
# tools, or any other detail while trying to reformat at the same time.
# ==============================================================================

def get_ai_data(resume_text, status_callback=None):

    # ------------------------------------------------------------------
    # PASS 1 — Extract everything, nothing filtered, nothing shortened
    # ------------------------------------------------------------------
    msg = "Pass 1: Extracting raw data from resume..."
    print(msg)
    if status_callback: status_callback(msg)

    pass1_system = """
    You are a meticulous resume data extractor.
    Your ONLY job: extract EVERY piece of information from the resume — verbatim.
    Nothing filtered. Nothing summarized. Nothing omitted.

    RULES:
    - **CRITICAL**: Distinguish between skills listed in a dedicated "Skills", "Technical Skills", or "Core Competencies" section vs. skills mentioned in work experience.
    - Extract ALL work experience entries with their FULL original text.
      Do NOT shorten, paraphrase, merge, or skip any entry or bullet.
    - Copy the professional summary / objective EXACTLY word for word.
    - Extract ALL certifications (including in-progress), awards, recognitions,
      client appreciations, and commendations.
    - Extract ALL education entries with degree, institution, year, score.
    - Preserve exact dates, numbers, percentages, and metrics as written.
    - NEVER truncate lists with "etc.", "and others", or similar.
    - If a field is genuinely absent, use "" or []. Never use null.
    """

    pass1_user = f"""
Resume text:
{resume_text}

Return this JSON — extract every detail exactly as it appears:
{{
  "personal_info": {{
    "name": "",
    "email": "",
    "phone": "",
    "location": ""
  }},
  "raw_summary": "Copy the professional summary/objective WORD FOR WORD as written. Do not change a single word.",
  "education": [
    {{"degree": "", "institution": "", "year": "", "score": ""}}
  ],
  "explicit_skills_section": {{
    "section_name": "Exact name of the skills section (e.g. 'Technical Skills', 'Core Competencies')",
    "content": "Copy the content of this section EXACTLY as written. If it's a list, keep it as a list."
  }},
  "experience_inferred_skills": {{
    "programming_languages": [],
    "frontend_frameworks": [],
    "backend_frameworks": [],
    "databases": [],
    "cloud_devops": [],
    "testing": [],
    "tools_ides": [],
    "css_styling": [],
    "version_control": [],
    "methodologies": [],
    "soft_skills": [],
    "domain_knowledge": [],
    "other": []
  }},
  "certifications": [],
  "awards_recognition": [],
  "all_experience": [
    {{
      "project_name": "Named project/product/system exactly as written",
      "organisation": "Employer or client company name",
      "duration": "Exact date range as written e.g. Sep 2021 - Present",
      "role": "Exact job title",
      "project_description": "Copy the project description EXACTLY as written in the resume",
      "responsibilities": ["Copy each responsibility bullet EXACTLY as written — do not paraphrase"],
      "tools_used": ["tool names only, no descriptions"],
      "technologies_used": ["technology and language names only"],
      "skills_used": ["skill names only"],
      "achievements": ["Copy each achievement EXACTLY as written with all numbers/percentages"],
      "client_appreciation": ["Copy any client quotes or recognition exactly as written"]
    }}
  ]
}}
"""

    pass1_response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": pass1_system},
            {"role": "user",   "content": pass1_user}
        ],
        response_format={"type": "json_object"}
    )
    raw_data = json.loads(pass1_response.choices[0].message.content)
    print("✓ Pass 1 complete.")
    
    # ------------------------------------------------------------------
    # PASS 2 — Format into exact template structure
    # ------------------------------------------------------------------
    msg = "Pass 2: Formatting into Zensar template structure..."
    print(msg)
    if status_callback: status_callback(msg)

    pass2_system = """
    You are a senior HR specialist at Zensar Technologies formatting a resume
    into Zensar's standard Word template. You will receive fully extracted raw
    data and must produce a structured JSON that feeds directly into the template.

    ══════════════════════════════════════════════════════════════════════
    FIELD INSTRUCTIONS
    ══════════════════════════════════════════════════════════════════════

    summary:
      Copy raw_summary EXACTLY as written in the resume. Do not rewrite,
      expand, or paraphrase. Preserve the candidate's own words.
      Only write a new summary if raw_summary is completely empty.

    education[]:
      Include ALL education entries (including school, intermediate, and college degrees).
      Each entry MUST include the full qualification name, institution, and score.
      Each entry: degree (string), year (string).
      Set 'degree' to the full text of the degree/qualification, institution, and score (e.g. "B.Tech (CSE) – CMR College, JNTUH | 78%").
      Set 'year' to the year/duration (e.g. "2016").

    skills[]:
      **CRITICAL**: Populate this list using ONLY the content from "explicit_skills_section".
      Do NOT include skills that are ONLY found in "experience_inferred_skills" unless the explicit section is completely empty.
      
      Group the skills from the explicit section into meaningful categories.
      If the candidate already grouped them (e.g., "Languages", "Frameworks"), PRESERVE those groups.
      If not, categorize them using standard industry terms:
        "Programming Languages", "Frontend Frameworks", "Backend & APIs",
        "Databases", "Cloud & DevOps", "Testing & QA", "Tools & IDEs",
        "CSS & Styling", "Version Control", "Methodologies", etc.
      
      primary:   most important / most used skills in that category (comma-separated string)
      secondary: less frequent skills in that category (comma-separated string)
      If secondary is empty, use "".

    expertise_areas[]:
      2-3 high-level domain strengths only, e.g.:
      "Full Stack Web Development", "Cloud & DevOps", "Agile Delivery".

    certifications[]: plain list of certification name strings.
    awards[]:         plain list of award/recognition strings.

    experience_summary[]:
      ── This is the MOST CRITICAL field ──
      The Word template loops over this array TWICE:
        Once as "s"   → fills the short Experience Summary table
        Once as "exp" → fills the detailed Professional Experience section
      Therefore every object MUST contain ALL keys listed below.
      Include ALL work experience entries. Most recent first. Do not skip any.

      SHORT TABLE KEYS (used by the Experience Summary table):
        role         → Exact job title string
        years        → Date range string, e.g. "Sep 2021 - Present"
        skills       → 5-8 key skills for this role as a comma-separated STRING
                       (not a list — a single string like "Angular, Java, AWS, Git")
        achievements → List of 2-3 SHORT one-line bullet strings for this role.
                       Keep each to one line. Use exact numbers/% from the resume.

      DETAILED SECTION KEYS (used by the Professional Experience section):
        project_name        → Named project/product/system. If none found, use org name.
        organisation        → Employer or client company name.
        duration            → Exact date range string (same value as years above).
        project_description → Copy the project description from the resume AS-IS.
                              If absent, write 2 clear lines about what the project does.
        key_activities      → Copy responsibility bullets from the resume AS-IS.
                              If vague, rewrite actively: "Developed...", "Led...", "Built..."
                              4-6 bullets. Each must describe one concrete specific action.
                              NEVER write "collaborated with team" without specifics.
        tools_used          → List of tool name strings only. No descriptions.
                              e.g. ["JIRA", "Git", "Postman", "VS Code", "Jenkins"]
        technologies_used   → List of technology/language name strings only.
                              e.g. ["Java", "Angular", "PostgreSQL", "Spring Boot"]
        skills_used         → List of skill name strings only.
                              e.g. ["Agile", "Code Review", "REST API Design"]
        major_achievements  → List of quantified achievement strings.
                              Use exact numbers/percentages from the resume.
                              If no numbers, describe the outcome specifically.
        client_appreciation → List of client quote/recognition strings.
                              Use [] if none found.

    ══════════════════════════════════════════════════════════════════════
    GENERAL RULES
    ══════════════════════════════════════════════════════════════════════
    - Fill EVERY field. Use "" for missing strings, [] for missing lists.
    - NEVER return null or None for any field.
    - Preserve all dates, numbers, and percentages exactly as extracted.
    - Do NOT invent data not present in the raw extraction.
    - Do NOT merge separate jobs into one entry.
    """

    pass2_user = f"""
Here is the fully extracted raw resume data:
{json.dumps(raw_data, indent=2)}

Produce the final JSON. Return ONLY this structure with no extra text:
{{
  "personal_info": {{
    "name": "",
    "email": "",
    "phone": "",
    "location": ""
  }},
  "summary": "Copied verbatim from resume",
  "education": [
    {{"degree": "", "year": ""}}
  ],
  "skills": [
    {{
      "category": "Category Name",
      "primary": "Skill A, Skill B, Skill C",
      "secondary": "Skill D, Skill E"
    }}
  ],
  "expertise_areas": ["Area 1", "Area 2"],
  "certifications": [],
  "awards": [],
  "experience_summary": [
    {{
      "role":                 "Exact job title",
      "years":                "Sep 2021 - Present",
      "skills":               "Skill A, Skill B, Skill C, Skill D, Skill E",
      "achievements":         ["Short achievement 1 with numbers", "Short achievement 2"],
      "project_name":         "Named project or system",
      "organisation":         "Company name",
      "duration":             "Sep 2021 - Present",
      "project_description":  "2-3 line description of the project",
      "key_activities":       ["Developed...", "Led...", "Implemented..."],
      "tools_used":           ["JIRA", "Git", "Postman"],
      "technologies_used":    ["Java", "Angular", "PostgreSQL"],
      "skills_used":          ["Agile", "Code Review", "REST API Design"],
      "major_achievements":   ["Reduced X by 30%", "Led team of 6"],
      "client_appreciation":  []
    }}
  ]
}}
"""

    pass2_response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": pass2_system},
            {"role": "user",   "content": pass2_user}
        ],
        response_format={"type": "json_object"}
    )
    final_data = json.loads(pass2_response.choices[0].message.content)
    print("✓ Pass 2 complete.")
    return final_data


# ==============================================================================
# STEP 3: SANITIZE
# Ensures every template variable exists and is never None.
# Keys verified against ravan.docx XML — do not rename them.
# ==============================================================================

def sanitize_data(data):

    # Personal info
    data.setdefault("personal_info", {})
    for key in ["name", "email", "phone", "location"]:
        data["personal_info"].setdefault(key, "")
        if data["personal_info"][key] is None:
            data["personal_info"][key] = ""

    # Top-level strings
    data.setdefault("summary", "")
    if data["summary"] is None:
        data["summary"] = ""

    # Top-level lists
    for key in ["education", "skills", "expertise_areas",
                "certifications", "awards", "experience_summary"]:
        data.setdefault(key, [])
        if data[key] is None:
            data[key] = []

    # Education entries  →  e.degree, e.year
    for edu in data["education"]:
        for key in ["degree", "year"]:
            edu.setdefault(key, "")
            if edu[key] is None:
                edu[key] = ""

    # Skill entries  →  handled by populate_skills_table()
    for skill in data["skills"]:
        for key in ["category", "primary", "secondary"]:
            skill.setdefault(key, "")
            if skill[key] is None:
                skill[key] = ""

    # experience_summary entries — ALL keys from BOTH template sections
    # Verified from ravan.docx XML: lines 1183-1933
    exp_defaults = {
        # Section 1 keys  (alias: s)
        "role":                 "",
        "years":                "",
        "skills":               "",
        "achievements":         [],
        # Section 2 keys  (alias: exp)
        "project_name":         "",
        "organisation":         "",
        "duration":             "",
        "project_description":  "",
        "key_activities":       [],
        "tools_used":           [],
        "technologies_used":    [],
        "skills_used":          [],
        "major_achievements":   [],
        "client_appreciation":  [],
    }
    for entry in data["experience_summary"]:
        for key, default in exp_defaults.items():
            entry.setdefault(key, default)
            if entry[key] is None:
                entry[key] = default

    return data


# ==============================================================================
# REMOVE MARKDOWN FORMATTING
# ==============================================================================

def remove_markdown_formatting(text):
    """Removes all markdown formatting from text."""
    if not isinstance(text, str):
        return text
    
    # Remove bold/italic markers
    text = text.replace("**", "").replace("*", "").replace("__", "").replace("_", "")
    
    # Remove headers (#)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        cleaned_lines.append(line.lstrip('#').strip())
    
    return '\n'.join(cleaned_lines)

def clean_markdown_from_data(data):
    """Recursively removes markdown from all string values in the dictionary."""
    if isinstance(data, dict):
        for key, value in data.items():
            data[key] = clean_markdown_from_data(value)
    elif isinstance(data, list):
        for i in range(len(data)):
            data[i] = clean_markdown_from_data(data[i])
    elif isinstance(data, str):
        return remove_markdown_formatting(data)
    
    return data
    

# ==============================================================================
# STEP 4: TABLE BORDERS
# ==============================================================================

def add_table_borders(table):
    """Adds full grid borders (outer frame + all inner lines) to a Word table."""
    tbl   = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for xml_tag in ["w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"]:
        existing = tblBorders.find(qn(xml_tag))
        if existing is not None:
            tblBorders.remove(existing)
        border = OxmlElement(xml_tag)
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)


# ==============================================================================
# STEP 5: SKILLS TABLE POPULATION
# ==============================================================================

def find_table_by_header(doc, header_text):
    for table in doc.tables:
        first_row_text = ' '.join([cell.text for cell in table.rows[0].cells]).lower()
        if header_text.lower() in first_row_text:
            return table
    return None


def populate_skills_table(doc, skills_data, status_callback=None):
    """Finds the skills table and repopulates it with AI data + borders."""
    msg = "Populating skills table..."
    print(msg)
    # Status callback is optional here as it's a sub-step
    
    skills_table = find_table_by_header(doc, "category")
    if not skills_table:
        skills_table = find_table_by_header(doc, "primary")
    if not skills_table:
        print("⚠ Warning: Skills table not found. Skipping.")
        return

    num_columns = len(skills_table.rows[0].cells)

    # Remove all data rows, keep header only
    for i in range(len(skills_table.rows) - 1, 0, -1):
        skills_table._element.remove(skills_table.rows[i]._element)

    # Add fresh rows
    for i, skill in enumerate(skills_data, 1):
        row = skills_table.add_row()
        if num_columns >= 4:       # [No. | Category | Primary | Secondary]
            row.cells[0].text = str(i)
            row.cells[1].text = skill.get('category',  '')
            row.cells[2].text = skill.get('primary',   '')
            row.cells[3].text = skill.get('secondary', '')
        elif num_columns == 3:     # [Category | Primary | Secondary]
            row.cells[0].text = skill.get('category',  '')
            row.cells[1].text = skill.get('primary',   '')
            row.cells[2].text = skill.get('secondary', '')

    add_table_borders(skills_table)
    print(f"✓ Added {len(skills_data)} skill rows with borders.")


# ==============================================================================
# STEP 6: DOCUMENT GENERATION
# ==============================================================================

def generate_doc(data, template_file, output_file, status_callback=None):
    msg = "\nGenerating Word document..."
    print(msg)
    if status_callback: status_callback("Generating Word document...")

    data = sanitize_data(data)
    data = clean_markdown_from_data(data)
      # Enforce limits
    # if len(data['skills']) > 7:
    #     data['skills'] = data['skills'][:7]
    # if len(data['experience_summary']) > 5:
    #     data['experience_summary'] = data['experience_summary'][:5]

    # Render Jinja2 template
    doc = DocxTemplate(template_file)
    doc.render(data)
    doc.save(output_file)

    # Post-process: repopulate skills table with proper borders
    msg = "Post-processing: populating skills table..."
    print(msg)
    if status_callback: status_callback(msg)
    
    rendered_doc = Document(output_file)
    populate_skills_table(rendered_doc, data['skills'])
    rendered_doc.save(output_file)

    print(f"\n✅ Resume saved to: {output_file}")


# ==============================================================================
# MAIN
# ==============================================================================


def process_resume(pdf_path, template_path, output_path, status_callback=None):
    """
    Main function to process a resume PDF and generate a Word document.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file '{template_path}' not found.")
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file '{pdf_path}' not found.")

    if status_callback: status_callback("PDF Received...")
    
    raw_text = extract_text_from_pdf(pdf_path, status_callback)
    resume_data = get_ai_data(raw_text, status_callback)

    print("\n" + "=" * 80)
    print("FINAL EXTRACTED + FORMATTED DATA:")
    print("=" * 80)
    print(json.dumps(resume_data, indent=2))
    print("=" * 80 + "\n")

    generate_doc(resume_data, template_path, output_path, status_callback)
    
    if status_callback: status_callback("Done!")
    return output_path


# ==============================================================================
# MAIN
# ==============================================================================

if __name__ == "__main__":
    try:
        process_resume(PDF_PATH, TEMPLATE_PATH, OUTPUT_PATH)
    except Exception as e:
        print(f"Error: {e}")
